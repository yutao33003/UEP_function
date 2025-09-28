import os
import re
import time
import pdfplumber
import filetype
from googletrans import Translator
from docx import Document
import nltk
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')
from nltk.tokenize import sent_tokenize

def chunk_text(text, chunk_size=3000):
    sentences = sent_tokenize(text) # 分句    
    chunks = []
    current_chunk = ""
    
    for sentence in sentences:
        sentence_len = len(sentence)
        
        # 長句子處理：如果句子長度超過 chunk_size，則進一步分割
        if sentence_len > chunk_size:
            if current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = ""
            split_long_sentence(sentence, chunks)
            continue
        
        # 一般情況：是否加入當前 chunk 會超出限制
        if len(current_chunk) + sentence_len + 1 > chunk_size:
            chunks.append(current_chunk.strip())
            current_chunk = sentence + " "
        else:
            current_chunk += sentence + " "
    
    if current_chunk:
        chunks.append(current_chunk.strip())
    return chunks

def split_long_sentence(sentence, chunk, chunk_size = 3000):
    parts = re.split(r'([，；,;])', sentence)
    current = ""
    for part in parts:
        if len(current) + len(part) > chunk_size:
            if current:
                chunk.append(current)
            current = part
        else:
            current += part
    if current:
        chunk.append(current)

def translate_chunk(chunk, dest_lang, translator, max_retry=3, wait_secs = 2):
    num = 0
    while num < max_retry:
        try:
            translated = translator.translate(chunk, dest=dest_lang)
            return translated.text
        except Exception as e:
            print(f"translate_chunk error: {e}")
            if num < max_retry - 1:
                print(f"Retrying in {wait_secs} seconds...")
                time.sleep(wait_secs)
            else:
                print("Max retries reached. Skipping this chunk.")
                return chunk

        num += 1

def save_translated_text(translated_chunks, output_path):
    doc = Document()
    for line in translated_chunks:
        doc.add_paragraph(line)
    doc.save(output_path)

def deter_file_type(path):
    kind = filetype.guess(path)
    if kind is None:
        print("Cannot guess file type")
        return 
    else:
        type = kind.extension
        return type


def pdf_segmenter(path):
    all_text_list = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                all_text_list.append(page_text)
    return all_text_list

def docx_segmenter(path):
    doc = Document(path)
    all_text_list = []
    for page in doc.paragraphs:
        text = page.text.strip()
        if text:
            all_text_list.append(text)
    return all_text_list

def save_translated_to_word(translated_text, output_word_path):
    doc = Document()
    for line in translated_text.split("\n"):
        doc.add_paragraph(line)
    doc.save(output_word_path)

def export_file(dest_code, origin_file_path, dest_file_path, ):
    file_type = deter_file_type(origin_file_path)
    doc = Document()
    translator = Translator()
    all_text_list = []

    if file_type == "docx":
        all_text_list = docx_segmenter(origin_file_path)
    elif file_type == "pdf":
        all_text_list =pdf_segmenter(origin_file_path)
    else: # 假設是純文本文件
        with open(origin_file_path, "r", encoding="utf-8") as f:
            text = f.read()
            all_text_list = [text]

    for para in all_text_list:
        chunks = chunk_text(para)
        translate_paragraph = ""
        for chunk in chunks:
            translated_text = translate_chunk(chunk, dest_code, translator)
            translate_paragraph += translated_text
        doc.add_paragraph(translate_paragraph)

    base_name = os.path.splitext(os.path.basename(origin_file_path))[0]
    new_file_name = f"{base_name}_{dest_code}_traslate.docx"
    full_dest_path = os.path.join(dest_file_path, new_file_name)
    doc.save(full_dest_path)
    os.startfile(full_dest_path)
            
